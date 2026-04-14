from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Inches
from PIL import Image, ExifTags
import os

app = Flask(__name__, template_folder='templates')

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# -------- RUTA PRINCIPAL --------
@app.route('/')
def index():
    return render_template('index.html')

# -------- GENERAR WORD --------
@app.route('/generar', methods=['POST'])
def generar():
    files = request.files.getlist("imagenes")
    
    doc = Document()

    from docx.enum.table import WD_TABLE_ALIGNMENT

    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, file in enumerate(files):
        row = (i % 4) // 2
        col = (i % 4) % 2

        cell = table.rows[row].cells[col]

        # 📁 Guardar imagen
        path = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(path)

        # 🔥 CORREGIR ROTACIÓN
        image = Image.open(path)

        try:
            for orientation in ExifTags.TAGS.keys():
                if ExifTags.TAGS[orientation] == 'Orientation':
                    break

            exif = image._getexif()

            if exif is not None:
                orientation_value = exif.get(orientation)

                if orientation_value == 3:
                    image = image.rotate(180, expand=True)
                elif orientation_value == 6:
                    image = image.rotate(270, expand=True)
                elif orientation_value == 8:
                    image = image.rotate(90, expand=True)

        except:
            pass

        image.save(path)

        # 📏 AJUSTE AUTOMÁTICO DE TAMAÑO
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()

        run.add_picture(path, width=Inches(3))  # 🔥 Ajuste clave

        # 📄 Nueva página cada 4 imágenes
        if (i + 1) % 4 == 0 and i != 0:
            doc.add_page_break()
            table = doc.add_table(rows=2, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

    output = "resultado.docx"
    doc.save(output)
    
    return send_file(output, as_attachment=True)

# -------- EJECUTAR --------
if __name__ == '__main__':
    app.run(host='0.0.0.0', port=10000)